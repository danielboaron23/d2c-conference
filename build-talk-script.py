#!/usr/bin/env python3
"""
Build the Hebrew talk script for the Design AI 7.0 deck (index-v4.html).
Outputs a .docx file with one section per slide:
  - Slide marker + brief visual context (in English, dimmed)
  - Full Hebrew script (what to say, word-for-word)
  - Transition note to the next slide
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    """Mark a paragraph as RTL so Hebrew renders correctly in Word."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11, rtl=True, align=None):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    return p


def add_visual_note(doc, text):
    """Small grey italic note describing what's on screen (in English)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[ON SCREEN] " + text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def add_transition(doc, text):
    """Italic stage direction for transitions (Hebrew, RTL)."""
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run("« מעבר: " + text + " »")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x99, 0x55)


def add_speech(doc, text, size=12):
    """The actual speaker line — main script body. Hebrew, RTL."""
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)


def slide_header(doc, num, title_he, title_en, slide_id):
    h = doc.add_heading(f"שקף {num} · {title_he}", level=2)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x0A, 0x66, 0x2E)
    # English subtitle line under the heading
    p = doc.add_paragraph()
    r = p.add_run(f"data-slide=\"{slide_id}\"  ·  {title_en}")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ------------------------------------------------------------
# Build document
# ------------------------------------------------------------

doc = Document()

# Page margins — slightly wider for more comfortable reading on stage
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ─── Cover ───
title = doc.add_heading("Design AI 7.0 — תסריט הרצאה", level=0)
set_rtl(title)
for r in title.runs:
    r.font.name = "Arial"

subtitle = doc.add_paragraph()
set_rtl(subtitle)
sr = subtitle.add_run("Where Is Product Design Heading?  ·  דניאל בוארון  ·  March 2026")
sr.font.name = "Arial"
sr.font.size = Pt(13)
sr.italic = True
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_para(doc, "")
add_para(doc,
    "המסמך הזה הוא תסריט מלא — מה להגיד, סלייד אחרי סלייד. "
    "הטקסט בשחור הוא מה שתגיד; השורות הקטנות באפור הן הקשר חזותי או הוראות במה. "
    "אורך משוער: 22–25 דקות. אם אתה במחסור בזמן — קצרו את הקטעים המסומנים [אופציונלי] "
    "ב-act 4 (deep dive בתוך flow). אם יש לך שעון, סמנו ביעד 12 דקות מסיום ב-flow-bridge.",
    italic=True, color=RGBColor(0x55, 0x55, 0x55), size=10)

doc.add_page_break()

# ============================================================
# ACT 1 · HOOK (slides 1-2) — ~2 min
# ============================================================
add_heading(doc, "Act 1  ·  הוק פתיחה", level=1)
add_para(doc, "כ־2 דקות", italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 1 ───
slide_header(doc, 1, "פתיחה — איפה עיצוב מוצר הולך", "Opening Hook (GLSL Hills)", "1")
add_visual_note(doc, "Animated 3D hills, big title fades in: 'Where Is Product Design Heading?' · subtitle: 'The tools changed. The roles are changing. The next 20 minutes might change how you think about both.'")

add_speech(doc,
    "ערב טוב לכולם. אם אתם פה, אתם כבר יודעים שמשהו זז. אז במקום שאני אבוא ואגיד לכם "
    "ש־AI 'הולך לשנות את הענף' — בואו נתחיל אחרת. בואו נתחיל בשאלה.")

add_speech(doc,
    "איפה עיצוב המוצר הולך? לא בעוד עשור — בעוד שלוש שנים. זאת לא שאלה תאורטית. "
    "זה משהו שכל אחד מאיתנו פה יצטרך לענות עליו, מאוד בקרוב, מול מנהל שאומר לנו "
    "‘אנחנו צריכים שאתה תעשה גם את זה’.")

add_speech(doc,
    "ב-20 הדקות הקרובות אני אנסה לתת לכם תמונה ברורה. לא בועה, לא היפ, לא ‘העתיד מפחיד’. "
    "המציאות. מה כבר קורה בענף, מה אני עושה בעצמי כל יום, ואיך אני חושב שכל אחד מכם צריך "
    "להתכונן.")

add_transition(doc, "מעבירים לסלייד 2 — הצגה עצמית. ‘אני דניאל בוארון, ובוא נדבר רגע על מי אני’.")

# ─── SLIDE 2 ───
slide_header(doc, 2, "מי אני — Daniel Boaron", "Who Am I", "2")
add_visual_note(doc, "Photo on right. Left: Design AI logo, label 'Your Speaker', name 'Daniel Boaron', short bio, 3 stat cards (7 / 500+ / 15+), trusted-by marquee.")

add_speech(doc,
    "אני דניאל בוארון. אני מעצב, מפתח, ומייסד של Design AI — הפלטפורמה שמלמדת צוותי עיצוב "
    "בארץ איך לעבוד נכון עם AI. זאת הפעם השביעית שאנחנו מעלים את הכנס הזה, ולאורך הדרך "
    "אימנתי כבר מעל 500 מעצבים, ב-15 חברות שונות.")

add_speech(doc,
    "אתם רואים בתחתית הסלייד את הלוגואים של החברות שעבדתי איתן. כשאני אומר ‘עבדתי איתן’ — "
    "אני לא מתכוון לשליחת מצגת ל-Slack פעם בחצי שנה. אני מתכוון לשבת עם הצוותים האלה, "
    "להבין את הזרימה שלהם, ולבנות איתם את ה-AI Workflow שלהם. monday, Siemens, Intuit, "
    "Rubrik, Cellebrite, ועוד הרבה יותר.")

add_speech(doc,
    "זה לא קורות חיים — זה רק כדי שתבינו שכשאני מדבר על מה עובד ומה לא, אני מדבר ממקום של "
    "ניסיון אמיתי, של ראיתי הרבה צוותים נופלים על אותן טעויות, וראיתי גם מה באמת עובד.")

add_transition(doc, "מבריק לאקט הבא — ‘בוא נראה איפה הענף נמצא היום’.")

doc.add_page_break()

# ============================================================
# ACT 2 · THE CHANGE IS HERE (slides 3-3d) — ~4 min
# ============================================================
add_heading(doc, "Act 2  ·  השינוי כבר קורה", level=1)
add_para(doc, "כ־4 דקות  ·  3 → 3a → 3b → 3c → 3d", italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 3 ───
slide_header(doc, 3, "איפה השינוי הולך", "Where Are Changes Heading?", "3")
add_visual_note(doc, "Synth shader background. Big title centered: 'WHERE ARE THE CHANGES HEADING?'")

add_speech(doc,
    "אז בואו נחזור לשאלה. איפה השינוי הולך? כי לפעמים אנחנו טועים ומסתכלים על הקטנות — "
    "‘איזה כפתור AI חדש יצא בFigma השבוע?’ — אבל זה לא הסיפור. הסיפור הוא לאן הענף עצמו "
    "זז. ואני רוצה להראות לכם 4 שכבות. נתחיל עם אחת פשוטה: זה כבר קורה.")

add_transition(doc, "ללחוץ קליק. הסטריפ של החברות הגדולות מופיע.")

# ─── SLIDE 3a ───
slide_header(doc, 1, "זה כבר קורה", "This Is Already Happening", "3a")
add_visual_note(doc, "Center title 'THIS IS ALREADY HAPPENING' + horizontal logo strip: Figma · Vercel · Cursor · Airbnb · Shopify · Stripe · Linear")

add_speech(doc,
    "החברות שאתם רואים על המסך — Figma, Vercel, Cursor, Airbnb, Shopify, Stripe, Linear — "
    "אלה לא חברות שעושות AI. אלה החברות שלך — שאתה משתמש בהן יום-יום. וכל אחת מהן, ברמה "
    "של C-level, החליטה ש-AI הוא חלק מהמוצר. לא אופציונלי. לא ‘בקרוב’. עכשיו.")

add_speech(doc,
    "אז כשמישהו בא ואומר לי ‘עוד מוקדם בשבילי לעבוד עם AI’ — אני שואל אותו: יותר מוקדם "
    "ממי? כי החברות שמעצבות לך את החיים כבר עברו את הקו הזה.")

add_transition(doc, "מעבר לסלייד 3b — מספרים.")

# ─── SLIDE 3b ───
slide_header(doc, 1, "במספרים", "By The Numbers", "3b")
add_visual_note(doc, "Three big stat cards: 85% (designers say AI essential — Figma 2025) · 46% (code on GitHub written with AI — 2025) · $14B (Anthropic ARR, 10× in 3 years).")

add_speech(doc,
    "ועכשיו במספרים. שלושה נתונים שאני רוצה שתסתכלו עליהם.")

add_speech(doc,
    "ראשון: 85 אחוז ממעצבים ומפתחים אומרים שכישורי AI חיוניים לעתיד שלהם. זה מהדוח של "
    "Figma מ-2025. שמונים וחמישה אחוז. לא רוב — כמעט כולם.")

add_speech(doc,
    "שני: 46 אחוז מהקוד שעולה עכשיו ל-GitHub נכתב בעזרת AI. כמעט חצי. לא בסטארט-אפים. "
    "ב-GitHub הגלובלי. אם אתה מעצב ואתה עובד מול דבלופרים — מחצית מהקוד שלהם כבר נכתב "
    "ככה.")

add_speech(doc,
    "ושלישי, אולי הכי דרמטי: Anthropic — היצרן של Claude — הגיעה ב-Q1 השנה ל-14 מיליארד "
    "דולר ARR. צמיחה של פי 10 בשלוש שנים. זה לא היפ. זה כסף שמשולם על ידי חברות, "
    "כל חודש, כי המוצר עובד.")

add_transition(doc, "‘מי שאומרים את זה?’ — מעבר לסלייד 3c.")

# ─── SLIDE 3c ───
slide_header(doc, 1, "מה הענף אומר", "The Industry Speaks", "3c")
add_visual_note(doc, "Three quote cards: Brian Chesky (Airbnb) · Guillermo Rauch (Vercel) · Dylan Field (Figma).")

add_speech(doc,
    "ויש את האנשים שאומרים את זה. שלושה ציטוטים, משלושה מנכ”לים שאתם מכירים.")

add_speech(doc,
    "ברייאן צ’סקי, מנכ”ל Airbnb, אומר: ‘אם מעצבים לא יחבקו AI, העולם יעוצב בלעדיהם.’ "
    "תקראו את זה שוב — בלעדיהם. זה לא ‘אתם תפסידו טרנד’. זה ‘העולם ימשיך, אתם תתעצרו’.")

add_speech(doc,
    "גיירמו ראוך, מנכ”ל Vercel, אומר: ‘כולם הם מהנדסים עכשיו’. כל אחד שיוצר מוצר — מעצב, "
    "מנהל מוצר, אפילו שיווק — יש לו עכשיו כלים שמאפשרים לו לבנות. זה משנה את ההגדרה של "
    "‘מעצב’.")

add_speech(doc,
    "ודילן פילד, מנכ”ל Figma, אומר את הצד החיובי: ‘לעיצוב יש הזדמנות בעידן ה-AI להבדיל "
    "מוצרים’. זה הצד שאני הכי אוהב — בעידן שבו כולם יכולים לבנות מהר, מי שיודע לעצב יודע "
    "לעשות את ההבדל.")

add_transition(doc, "‘ומה זה אומר על השוק עצמו? מי מקבל עבודה?’ — מעבר ל-3d.")

# ─── SLIDE 3d ───
slide_header(doc, 1, "השליפה של HR", "The New Hiring Bar", "3d")
add_visual_note(doc, "Header 'WHAT COMPANIES WANT — THE NEW HIRING BAR' + four cards: 73% / 68% / $294K / 96%")

add_speech(doc,
    "וזה אולי הכי רלוונטי לקהל פה. השוק עצמו — ה-HR.")

add_speech(doc,
    "73 אחוז ממנהלי גיוס היום דורשים שליטה בכלי AI ממעצבים. דורשים. לא ‘מעדיפים’. "
    "זה כבר תנאי כניסה. לפני שלוש שנים זה לא היה שום שורה בתיאור משרה.")

add_speech(doc,
    "68 אחוז מהמעצבים עצמם — מסקר של חצי אלף איש — אומרים שהם חושבים שהם צריכים לדעת "
    "לקודד. לא ‘יודעים לקודד’. ‘חושבים שצריכים’. זאת תחושה משותפת בכל הענף.")

add_speech(doc,
    "294 אלף דולר — זה השכר העליון ל-Design Engineer ב-Vercel. תפקיד שלא היה קיים לפני "
    "שלוש שנים. כן, יש פה אימפליציטה: שילוב של עיצוב וקוד מקבל עכשיו את החבילה הכי גבוהה "
    "בענף.")

add_speech(doc,
    "ו-96 אחוז — והמספר הזה מטריד אותי הכי הרבה — מהמעצבים שכן למדו AI, למדו את זה לבד. "
    "לא דרך החברה שלהם, לא דרך אקדמיה. לבד. אז כשאני שואל ‘מי אחראי על ההכשרה שלך?’ "
    "התשובה היום היא ‘אתה’.")

add_transition(doc, "סוגרים את אקט המספרים. ‘אז יש שינוי. אבל השינוי הוא לא רק בכלים — הוא בתפקיד עצמו’.")

doc.add_page_break()

# ============================================================
# ACT 3 · THE NEW ROLE (slides 6a-6b, 4-4c) — ~4 min
# ============================================================
add_heading(doc, "Act 3  ·  התפקיד מתפרק ומתאחד", level=1)
add_para(doc, "כ־4 דקות  ·  6a → 6b → 4 → 4e → 4c", italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 6a ───
slide_header(doc, 1, "תהליך העבודה — המשולש", "The Work Process — Triangle", "6a")
add_visual_note(doc, "Big triangle with FE / PM / PD nodes. Dashed lines connect them. A dot orbits between roles. Dissolve button: 'What if boundaries blur?'.")

add_speech(doc,
    "ככה תהליך עבודה נראה היום ב-95 אחוז מהארגונים. שלושה תפקידים. PM כותב מסמך, PD "
    "מעצב מסך, FE כותב קוד. הכדור עובר ביניהם — תיקונים, מסמכים, מצגות, פגישות יישור.")

add_speech(doc,
    "וזה עובד. עובד מצוין כשהמוצר ידוע, השוק יציב, וכל אחד יודע מה הוא עושה. אבל יש "
    "בעיה אחת — זה איטי. כי כל פעם שהכדור עובר בין אנשים, הוא מאבד תנופה.")

add_speech(doc,
    "אז שואלים את עצמינו — מה אם הגבולות יטשטשו?")

add_transition(doc, "מורפינג ל-6b — שלוש העיגולים מתאחדים לעיגול אחד.")

# ─── SLIDE 6b ───
slide_header(doc, 1, "מעגל היוצר", "The Maker Circle", "6b")
add_visual_note(doc, "Three role-circles dissolve into one big circle: 'The Product Builder · FE · PM · PD'. Subtitle: 'Boundaries dissolved. One person, all the skills.'")

add_speech(doc,
    "זה אדם אחד. ה-Product Builder. הוא יודע מוצר, הוא יודע עיצוב, הוא יודע פיתוח. "
    "אז שני דברים חשוב להגיד פה. אחד — זה לא חדש. סטיב ג'ובס היה כזה. ג'ון איב היה כזה. "
    "מייסדי כל הסטארט-אפ הראשונים שאתם מכירים היו כאלה.")

add_speech(doc,
    "מה שחדש זה שעכשיו, עם AI, גם אתם יכולים להיות כאלה. אתם לא צריכים להפוך לדבלופרים. "
    "אתם פשוט יכולים לעבוד עם כלים שמאפשרים לכם לעצב ולבנות באותו זמן. אתם המוח, "
    "ה-AI הוא הידיים.")

add_speech(doc,
    "אבל לפני שאני נכנס לאיך — אנחנו חייבים לשאול את השאלה ההיא, שכולם שואלים, אבל אף "
    "אחד לא רוצה לענות עליה ישר.")

add_transition(doc, "מעבר טייפרייטר — קולות גשם, השאלה צצה אות אחרי אות.")

# ─── SLIDE 4 ───
slide_header(doc, 1, "מעצבים צריכים…?", "Should Designers …?", "4")
add_visual_note(doc, "Rain canvas. Big typewriter line: 'SHOULD DESIGNERS [CODE / SHIP / PUSH / OPEN PRS / BUILD] ?' — words cycle with typewriter effect.")

add_speech(doc,
    "מעצבים צריכים לקודד? ב-90 הדקות הראשונות בכל הרצאה אני שומע את השאלה הזאת. ובעצם, "
    "זאת לא השאלה. או יותר נכון — זאת השאלה הלא נכונה.")

add_speech(doc,
    "כי השאלה היא לא ‘האם מעצבים צריכים לקודד’. השאלה היא ‘האם מעצבים צריכים לבנות’. "
    "‘האם מעצבים צריכים לסגור את הלולאה’. ‘האם מעצבים צריכים לפתוח PR’. שימו לב — "
    "אני לא אומר ‘לכתוב קוד’. אני אומר ‘לבנות’. הכלים השתנו.")

add_transition(doc, "מעבר לוידאו של 4e — מראים את הלולאה: עיצוב ⇄ קוד.")

# ─── SLIDE 4e ───
slide_header(doc, 1, "הלולאה — קוד ⇄ עיצוב", "The Loop — Code ↔ Figma", "4e")
add_visual_note(doc, "Full-bleed video showing live: a designer types in Cursor → code generates → Figma updates → designer adjusts in Figma → code rebuilds. Badge: 'The Workflow'.")

add_speech(doc,
    "מה שאתם רואים פה זה אני, יום יום. שמאל — Cursor, ימין — Figma. אני כותב משהו בעיצוב, "
    "ה-AI כותב את הקוד. אני משנה משהו בקוד, Figma מתעדכן. הלולאה הזאת — קוד וצ’יפמה "
    "מתחברים — היא הכלי החזק ביותר שיש לי כיום.")

add_speech(doc,
    "זה לא ‘המעצב הופך לדבלופר’. זה ‘המעצב סגר את הלולאה’. הוא רואה את התוצאה האמיתית — "
    "לא mockup ב-Figma — וקובע אם זה עובד או לא. ובלי שיצטרך לחכות שבועיים שדבלופר יבנה.")

add_transition(doc, "‘אז יש כלי. רגע, יש כלי? יש 50 כלים. בוא נדבר על זה’.")

# ─── SLIDE 4c ───
slide_header(doc, 1, "האקוסיסטם — הרבה כלים", "The Toolkit — So Many Tools", "4c")
add_visual_note(doc, "Header: 'The Ecosystem · So Many Tools. Where Do You Start?' + 16 logos floating in a kinetic stage (Figma, Cursor, Claude, ChatGPT, Gemini, Copilot, etc.).")

add_speech(doc,
    "וזה הצרה האמיתית. תסתכלו על הסלייד הזה. כל לוגו פה הוא כלי AI שמיועד למעצבים או "
    "למפתחים. בכל שבוע יוצא חדש. ובכל שבוע יש לינקדאין שמישהו מודיע ש-‘זה הכלי שמשנה הכל’.")

add_speech(doc,
    "אז שלוש דקות אחרי שראיתי את השלישי, אני באמת מבולבל. הראשון אמר ‘זה משנה הכל’. השני "
    "אמר ‘אבל בא חדש’. השלישי אמר ‘שכחו מה שאמרתי, יש משהו טוב יותר’. מה אני אמור לעשות?")

add_speech(doc,
    "התשובה — והיא הולכת להפתיע אתכם — היא לא לנסות את כולם. התשובה היא הפוכה.")

add_transition(doc, "מעבר ל-4d — Go Deep, Not Wide.")

doc.add_page_break()

# ============================================================
# ACT 4 · GO DEEP NOT WIDE (slide 4d) — ~1.5 min
# ============================================================
add_heading(doc, "Act 4  ·  לעומק, לא לרוחב", level=1)
add_para(doc, "כ־1.5 דקות  ·  4d", italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 4d ───
slide_header(doc, 1, "לעומק, לא לרוחב — הכלים שלי", "Go Deep, Not Wide", "4d")
add_visual_note(doc, "Center: 'GO DEEP, NOT WIDE'. Below: 4 tool cards (Figma · Cursor · Claude · GitHub). Footer badge: 'My Toolkit · The New Designer'.")

add_speech(doc,
    "ככה זה עובד אצלי. ארבעה כלים. רק ארבעה. Figma — לעיצוב. Cursor — סביבת הפיתוח. "
    "Claude — מנוע ה-AI שמדבר עם הכל. GitHub — לאן הקוד הולך בסוף.")

add_speech(doc,
    "כל יום אני נתקל במישהו ששואל אותי ‘למה אתה לא משתמש ב-X?’ או ‘ניסית את Y?’. "
    "התשובה תמיד אותה תשובה: כי אני שולט בארבעה האלה. שולט באמת. אני יודע לעשות איתם "
    "דברים שמישהו שניסה אותם 20 דקות לא יודע לעשות.")

add_speech(doc,
    "בנייה של מומחיות — depth — היא מה שעושה את ההבדל. רוחב — width — רק מבלבל אותך. "
    "אז קח את הכלים שלך, צמצם לרבעוני, וצלול עמוק.")

add_transition(doc, "‘ועכשיו — בוא אני אראה לכם איך נראה ה-flow שלי, מהתחלה עד הסוף’.")

doc.add_page_break()

# ============================================================
# ACT 5 · MY WORKFLOW (4-flow-master through 4-flow-recap) — ~10 min
# ============================================================
add_heading(doc, "Act 5  ·  הזרימה שלי, סוף-לסוף", level=1)
add_para(doc, "כ־10 דקות  ·  זה החלק העמוק. אם נגמר זמן — קצרו אחד מ-flow-1..5 לחצי דקה.",
    italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 4-flow-master ───
slide_header(doc, 1, "ה-Workflow שלי — 5 שלבים", "MY WORKFLOW (n8n-style)", "4-flow-master")
add_visual_note(doc, "n8n-style 5-node diagram (Discovery · Ideation/Design · UI Fix branch · Dev Prep · Ship). Cards reveal phase by phase.")

add_speech(doc,
    "זה ה-flow שלי. חמישה שלבים, מהרגע שמשהו נכנס ל-Jira, עד הרגע שזה ב-production.")

add_speech(doc,
    "Stage 1: Discovery. אני מבין מה אני בונה ולמה.")

add_speech(doc,
    "Stage 2: Ideation ועיצוב. אני בודק רעיונות, מעצב מסכים, מקבל החלטות.")

add_speech(doc,
    "Stage 3: זה ה-branch — אופציונלי. אם צריך תיקוני UI, אני נכנס ל-Figma, מסדר, חוזר.")

add_speech(doc,
    "Stage 4: Dev Prep / Implement. כאן הקוד נכתב.")

add_speech(doc,
    "Stage 5: Ship. PR יוצא, ה-feature ב-production.")

add_speech(doc,
    "וזה כולו — באני אחד. עם הכלים שראיתם.")

add_transition(doc, "‘שניה לפני שאני נכנס לעומק כל שלב — חשוב שתבינו שזה לא תמיד לינארי’.")

# ─── SLIDE 4-flow-nonlinear ───
slide_header(doc, 1, "זה לא חייב להיות לינארי", "It Doesn't Have To Be Linear", "4-flow-nonlinear")
add_visual_note(doc, "Top rail: minimized master flow. Title 'IT DOESN'T HAVE TO BE LINEAR · The flow adapts to the task — sometimes you start in Figma'. Bottom rail reveals an alt flow starting from Figma.")

add_speech(doc,
    "כי לפעמים, אני לא מתחיל ב-Jira. לפעמים אני מתחיל ב-Figma. למשל — מנהל המוצר שלי "
    "פתאום שולח לי mock, אומר ‘תראה לי איך זה נראה ב-prod’. אני לא צריך לכתוב מסמך. "
    "אני פותח Figma, מתקן, יורד ל-Cursor, וזה עובד.")

add_speech(doc,
    "ה-flow מסתגל למשימה. ולפעמים זה נכנס באמצע — אני מתחיל מ-Stage 3, חוזר ל-Stage 2, "
    "ויוצא ב-Stage 5. החיים האמיתיים. וכולם תקפים.")

add_transition(doc, "‘ועכשיו, בוא נראה את כל זה בפעולה. כל שלב, וידאו אמיתי’.")

# ─── SLIDE 4-flow-bridge ───
slide_header(doc, 1, "בוא נראה את זה בפעולה", "Let's See It In Action", "4-flow-bridge")
add_visual_note(doc, "Master flow rail morphs back to center. Header reveals: 'LET'S SEE IT IN ACTION — Stage by stage, real videos from my workflow'.")

add_speech(doc,
    "החלק הזה הוא הלב של ההרצאה. עכשיו, כל שלב — אני אראה לכם וידאו אמיתי, מהמסך שלי, "
    "מהיומיום. לא הדגמה. אמיתי.")

add_transition(doc, "מעבר ל-Stage 1.")

# ─── SLIDE 4-flow-1 ───
slide_header(doc, 1, "Stage 01 · Discovery", "Discovery — Claude Code in Cursor", "4-flow-1")
add_visual_note(doc, "Stage 01/05 header. Left: video player. Right: Discovery stack — Jira, GitHub, Claude Code (in Cursor).")

add_speech(doc,
    "Stage 1 — Discovery. הרגע שמשהו נכנס. בעיניי, זה הרגע שמכתיב את כל מה שיקרה אחר כך.")

add_speech(doc,
    "מה שאתם רואים פה זה Claude Code שיושב לי בתוך Cursor. אני לוחץ על קליק — והוא ניגש "
    "ל-Jira דרך MCP, מושך לי את הטיקט. מסתכל על GitHub, מבין מה כבר קיים בקוד, מה "
    "ה-pattern של הצוות. כל זה לפני שאני בכלל פתחתי Figma.")

add_speech(doc,
    "ולמה זה משנה? כי לפני שלוש שנים, הייתי מבזבז את שעת ה-discovery הראשונה בלקרוא "
    "את הטיקט, להבין הקשר, לחפש בוידאו פגישות. עכשיו ה-AI מסכם לי את כל זה בדקה. "
    "אני נכנס לעבודה כשאני כבר ‘בעניינים’.")

add_transition(doc, "ללחוץ קליק נוסף — הוידאו עולה למסך גדול. נותנים לו לרוץ ולספק.")

# ─── SLIDE 4-flow-2 ───
slide_header(doc, 1, "Stage 02 · Ideation / Design", "Ideation — Skills (PM + Designer)", "4-flow-2")
add_visual_note(doc, "Stage 02/05. Left: video. Right: two Claude Skill cards reveal one by one — 'Product Manager Skill (Lenny)' and 'Product Designer Skill (UI/UX Pro / Design System)'.")

add_speech(doc,
    "Stage 2 — Ideation. עכשיו אני צריך לחשוב.")

add_speech(doc,
    "ופה הסוד שלי. אני לא חושב לבד. אני בניתי שני Claude Skills שעובדים איתי. אחד אני "
    "קורא לו ‘Product Manager Skill’ — מבוסס על תפיסות של לני רכיצקי, מהפודקאסט שלו. "
    "כשאני מתלבט על משהו מוצרי, אני שואל את ה-Skill הזה: מה התעדוף? מה מטריצת ההשפעה?")

add_speech(doc,
    "השני — ‘Product Designer Skill’ — בנוי על UI/UX Pro וה-Design System של החברה שאני "
    "עובד איתה. הוא יודע איך לעצב לפי הסטנדרטים שלנו, לא של פינטרסט. הוא יודע מה ‘נכון’.")

add_speech(doc,
    "וזה השינוי האמיתי. ה-AI שלי הוא לא ‘ChatGPT גנרי’. הוא ה-AI שלי, עם הידע שלי, "
    "הסטנדרטים שלי, והגישה שלי.")

add_transition(doc, "‘וכשאני סוגר עיצוב — לפעמים אני צריך לתקן UI ספציפי. אז הולכים ל-branch’.")

# ─── SLIDE 4-flow-3 ───
slide_header(doc, 1, "Stage 03 · UI Fix (אופציונלי)", "UI Fix — Figma AI", "4-flow-3")
add_visual_note(doc, "Stage 03/05 BRANCH (orange). Left: video. Right: floating Figma MCP/Figma AI logo.")

add_speech(doc,
    "Stage 3 — אופציונלי. UI Fix. זה ה-branch הכתום שראיתם בדיאגרמה.")

add_speech(doc,
    "לפעמים, אני צריך לתקן משהו ויזואלי שנראה לא טוב. במקום לפתוח Figma ידני ולמשוך "
    "פיקסלים, אני משתמש ב-Figma AI. אני אומר לו ‘קח את ה-component הזה, התאם אותו "
    "לסטנדרט החדש’ — והוא עושה. לא תמיד מושלם — אבל מספיק טוב כדי שאני אשפר משם.")

add_speech(doc,
    "השלב הזה הוא ה-shortcut שלי. הוא חוסך לי שעות.")

add_transition(doc, "ללחוץ — הוידאו עולה למסך מלא. ‘אבל ברב המקרים, אני מדלג מ-Stage 2 ישר ל-Stage 4’.")

# ─── SLIDE 4-flow-4 ───
slide_header(doc, 1, "Stage 04 · Dev Prep / Implement", "Implementation — Code Review + Eng Lead", "4-flow-4")
add_visual_note(doc, "Stage 04/05 BUILD. Left: video (still poster, plays full-screen on click 3). Right: two Claude Skills reveal — 'Code Review' and 'Engineering Lead'.")

add_speech(doc,
    "Stage 4 — כאן הקוד נבנה. ופה אני משתמש בעוד שני Skills שאני יכול ממש להמליץ עליהם.")

add_speech(doc,
    "Code Review Skill — סקיל רשמי של Anthropic. הוא בוחן כל שינוי שאני עושה בקוד, "
    "מחפש בעיות, מבוטות, מציע שיפורים. זה כמו שיש לי senior engineer יושב לידי, "
    "מסתכל על הקוד שלי, ואומר ‘רגע, פה כדאי לחשוב שוב’.")

add_speech(doc,
    "Engineering Lead Skill — שני. הוא יודע את ה-codebase שלי, את ה-conventions של "
    "הצוות, את ה-PR style. הוא לא רק כותב קוד — הוא כותב קוד שעובר ב-PR בלי דרמה.")

add_speech(doc,
    "ובסוף, אני לוחץ עוד קליק — ורואים את הוידאו עצמו. תקדים — אני בקוד, מוסיף feature, "
    "ה-Skills עובדים ברקע.")

add_transition(doc, "‘ובסוף — Ship’.")

# ─── SLIDE 4-flow-5 ───
slide_header(doc, 1, "Stage 05 · Ship", "Ship — PR by parts", "4-flow-5")
add_visual_note(doc, "Stage 05/05 SHIP. Full-width video plays automatically. Single bullet card: 'PR by parts'.")

add_speech(doc,
    "Stage 5 — Ship. וזה הוידאו שאתם רואים עכשיו, רץ אוטומטית.")

add_speech(doc,
    "במקום PR ענק שאף אחד לא רוצה לסקור — אני שובר אותו לחתיכות. כל חתיכה היא PR קצר, "
    "ממוקד, שאפשר לסקור ב-5 דקות. ה-AI עוזר לי לחלק נכון. וה-codereviewer של הצוות שמח, "
    "כי הוא רואה PR קטנים ובהירים, לא monster של 2000 שורות.")

add_transition(doc, "‘וזה כל ה-flow. מהתחלה לסוף. בוא נראה את כולו פעם אחרונה’.")

# ─── SLIDE 4-flow-recap ───
slide_header(doc, 1, "זה ה-Flow", "That's The Flow — Recap", "4-flow-recap")
add_visual_note(doc, "Recap badge. Title 'THAT'S THE FLOW — From Jira ticket to shipped PR — every stage covered'. All 5 cards reveal in stagger.")

add_speech(doc,
    "זה הסיפור. מ-Jira, עד PR ב-production. חמישה שלבים, אדם אחד, ארבעה כלים. וכל שלב "
    "— Skills מותאמים שעובדים בשבילי.")

add_speech(doc,
    "תקראו לזה ‘flow’. אני קורא לזה ‘תפקיד חדש’.")

add_transition(doc, "‘אז בוא נחזור לשאלה. מעצבים צריכים לקודד? לא. אבל הם צריכים משהו אחר’.")

doc.add_page_break()

# ============================================================
# ACT 6 · WHAT'S NEXT (4d-1, 4d-adapt, 4d-focus, 4d-discipline) — ~3 min
# ============================================================
add_heading(doc, "Act 6  ·  להתאים את עצמך", level=1)
add_para(doc, "כ־3 דקות  ·  4d-1 → 4d-adapt → 4d-focus → 4d-discipline",
    italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 4d-1 ───
slide_header(doc, 1, "השאלה האמיתית", "Should Designers Be Builders / Makers?", "4d-1")
add_visual_note(doc, "Same typewriter style as slide 4. Cycles 'SHOULD DESIGNERS [CODE / BE BUILDERS / BE MAKERS / PUSH / OPEN PRS / SHIP] ?'")

add_speech(doc,
    "נחזור לשאלה ההיא. מעצבים צריכים לקודד? לא. הם צריכים להיות יוצרים. הם צריכים "
    "להיות בונים. הם צריכים לסגור את הלולאה.")

add_speech(doc,
    "‘לקודד’ זה הכלי. ‘לבנות’ זה ה-mindset.")

add_transition(doc, "מעבר ל-4d-adapt — מילים מתפרקות לחלקיקים, מתקבצות לכותרת חדשה.")

# ─── SLIDE 4d-adapt ───
slide_header(doc, 1, "מעצבים צריכים להתאים", "Designers Need to Adapt", "4d-adapt")
add_visual_note(doc, "Vapor particle effect — text 'DESIGNERS NEED TO ADAPT' forms from particles. Subtitle: 'The tools changed. The role is evolving.'")

add_speech(doc,
    "ולהיות יוצר — זה אומר להתאים. הכלים השתנו. התפקיד מתפתח. ואם אתה לא זז, אתה נשאר "
    "מאחור.")

add_speech(doc,
    "אז עכשיו השאלה היא: איך מתאימים? איפה משקיעים? כי, כמו שראיתם, יש 50 כלים. "
    "במה מתחילים?")

add_transition(doc, "‘אני אספר לכם איפה אני הייתי שם את הזמן שלי, אם הייתי מתחיל היום’.")

# ─── SLIDE 4d-focus ───
slide_header(doc, 1, "איפה הייתי משקיע — 3 הימורים", "Where I'd Focus Now — 3 Bets", "4d-focus")
add_visual_note(doc, "Header 'My Recommendation · WHERE I'D FOCUS NOW'. Three flip-cards: Bet 01 Skills (green) · Bet 02 Agents (purple) · Bet 03 Design.md (orange). Cards flip to reveal details.")

add_speech(doc,
    "שלושה הימורים. אם הייתי בקהל היום — אלה השלושה שהייתי שם עליהם.")

add_speech(doc,
    "ראשון: Skills. ה-AI שלך, עם ה-playbook שאתה כתבת. אל תחפש ‘ChatGPT הכי טוב’ — בנה "
    "Claude Skill עם הידע שלך. תאפיין אותו ל-workflow שלך. זה ה-leverage הכי גבוה שיש.")

add_speech(doc,
    "שני: Agents. צוות של agents שעובד במקביל בשבילך. אתה מנצח על תזמורת — הם מנגנים. "
    "Claude Managed Agents, OpenClaw — תכל’ס, זה החלק הכי מתקדם של AI היום. ולפי דעתי, "
    "תוך שנה זה יהיה standard.")

add_speech(doc,
    "שלישי: Design.md. קובץ Markdown אחד — שמכיל את ה-Design System שלך, בשפה שמכונות "
    "מבינות. כל agent שעובד אצלך קורא אותו. זה ה-source of truth החדש. Stitch של גוגל "
    "ו-Claude Design הולכים בדיוק לכיוון הזה.")

add_speech(doc,
    "אם אתם רוצים להיות מוכנים לשלוש שנים הבאות — אלה ההימורים שלי.")

add_transition(doc, "‘אבל יש דבר אחד אחרון, וזה אולי הכי חשוב’.")

# ─── SLIDE 4d-discipline ───
slide_header(doc, 1, "המשמעת — להוריד את ה-FOMO", "Drop the FOMO Manifesto", "4d-discipline")
add_visual_note(doc, "Three lines reveal in sequence: 'DROP THE FOMO.' / 'BUILD YOUR TOOLKIT.' / 'MASTER WHAT MATTERS.' Closing: 'What gives value — keep. What doesn't — let go.'")

add_speech(doc,
    "המשמעת. שלוש שורות.")

add_speech(doc,
    "ראשונה: Drop the FOMO. הורידו את הפחד שאתם מפסידים משהו. כל יום יש כלי חדש. כל יום "
    "מישהו אחר אומר ‘זה משנה הכל’. תפסיקו להקשיב.")

add_speech(doc,
    "שנייה: Build your toolkit. בנו את ערכת הכלים שלכם. ארבעה. חמישה. לא יותר. הכלים "
    "שעובדים בשבילכם.")

add_speech(doc,
    "שלישית: Master what matters. תהיו המאסטרים. depth, לא width. ככל שתהיו עמוקים "
    "יותר, ההשפעה שלכם תגדל.")

add_speech(doc,
    "מה שנותן ערך — תשמרו. מה שלא — תרפו.")

add_transition(doc, "‘זה השוק. זה התפקיד. וזה המסר שלי לכם’.")

doc.add_page_break()

# ============================================================
# ACT 7 · CLOSING (slides 10-12) — ~1 min
# ============================================================
add_heading(doc, "Act 7  ·  סגירה", level=1)
add_para(doc, "כ־1 דקה  ·  10 → 11 → 12", italic=True, color=RGBColor(0x77, 0x77, 0x77), size=10)

# ─── SLIDE 10 ───
slide_header(doc, 1, "מי שמתאים — מוביל", "Designers Who Adapt Lead", "10")
add_visual_note(doc, "Vapor particle effect — text 'DESIGNERS WHO ADAPT LEAD' forms from particles. Badge: 'The Proof'.")

add_speech(doc,
    "אז אם יש משהו אחד שתיקחו מההרצאה הזאת — שיהיה זה.")

add_speech(doc,
    "מעצבים שמתאימים — מובילים. לא מעצבים שיודעים לקודד הכי טוב. לא מעצבים שיודעים את "
    "כל הכלים. מעצבים שיודעים להתאים, להעמיק, ולסגור לולאות.")

add_transition(doc, "מעבר לסלייד תודות.")

# ─── SLIDE 11 ───
slide_header(doc, 1, "תודה — ותהנו", "Thank You & Have Fun", "11")
add_visual_note(doc, "Daniel's photo on right with badge 'Thank You — Design AI 7.0'. Left: 'Thank You & Have Fun.' + bio: 'Stay curious. Build the toolkit that fits you. Ship the work only you can ship.' + Design AI logo.")

add_speech(doc,
    "תודה רבה לכולם. אני מקווה שיצאתם עם משהו מעשי, לא רק היפ.")

add_speech(doc,
    "תהיו סקרנים. בנו את ערכת הכלים שמתאימה לכם. ושלחו את העבודה שרק אתם יכולים לשלוח.")

add_transition(doc, "מעבר לQR — אם רוצים להתחבר.")

# ─── SLIDE 12 ───
slide_header(doc, 1, "QR — שמרו על קשר", "Big QR — Scan to Connect", "12")
add_visual_note(doc, "Big QR code centered. Header: 'Let's keep in touch · SCAN TO CONNECT.' Caption: 'Daniel Boaron · Design AI'.")

add_speech(doc,
    "ואם רוצים להמשיך את השיחה — תסרקו את ה-QR. כל הקישורים — Linkedin, ה-newsletter, "
    "כל המשאבים שדיברתי עליהם היום, וגם הסליידים האלה — תמצאו שם. תודה!")

add_transition(doc, "מורידים את הראש, מחכים לשאלות.")

# ============================================================
# Appendix
# ============================================================
doc.add_page_break()
add_heading(doc, "נספח  ·  עצות לדוברים", level=1)

add_para(doc, "תזמון", bold=True, size=12)
add_para(doc, "•  בדקה 12 שלכם — אתם צריכים להיות באמצע ה-flow walkthrough. אם אתם עדיין על Act 3, קצרו.", size=11)
add_para(doc, "•  בדקה 18 — אתם צריכים להיות ב-‘Where I'd Focus Now’. אם עוד לא — דלגו על הוידאו של Stage 5.", size=11)
add_para(doc, "•  השאירו 2 דקות לתודות + שאלות.", size=11)

add_para(doc, "")
add_para(doc, "מה לעשות אם לא קורה משהו על המסך", bold=True, size=12)
add_para(doc, "•  כל סלייד עם וידאו — אם הוידאו לא עולה בקליק שלישי, יש כפתור Esc → רענן. דברו עליו: ‘בזמן שזה נטען...’.", size=11)
add_para(doc, "•  סלייד 4-flow-master ו-recap צריכים גם הם 5 קליקים בהדרגה — לא לחפז.", size=11)

add_para(doc, "")
add_para(doc, "מילים שיש להגיד נכון", bold=True, size=12)
add_para(doc, "•  ‘Claude Skills’ — לא ‘claude-skill-S’. הדגשה על ה-‘S’.", size=11)
add_para(doc, "•  ‘MCP’ — אותיות נפרדות, M-C-P, לא ‘mcp’ כמילה.", size=11)
add_para(doc, "•  ‘Cursor’ — לא ‘קוּרסור’. ‘Cursor’ באנגלית.", size=11)
add_para(doc, "•  שמות חברות לא ישראליות — תגידו אותם באנגלית. ‘monday.com’, ‘Vercel’.", size=11)

# Save
output = "/Users/danielboaron/d2c-conference/talk-script.docx"
doc.save(output)
print(f"OK · Saved: {output}")
